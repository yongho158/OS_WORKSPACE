from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "output"
REPORT_PATH = ROOT / "운영체제최종보고서_작성본.docx"
REVIEWED_REPORT_PATH = ROOT / "운영체제최종보고서_작성본_검토수정.docx"

KOREAN_FONT = "맑은 고딕"
CODE_FONT = "Consolas"
TEAM_NAME = "7팀"
MEMBER_NAME = "최용호"
MEMBER_ID = "2022152039"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        return list(csv.DictReader(file))


def find_template() -> Path | None:
    matches = sorted(ROOT.glob("운영체제최종보고서서식*.docx"))
    return matches[0] if matches else None


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = KOREAN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    normal.font.size = Pt(10)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = KOREAN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def set_run_font(run, font_name: str = KOREAN_FONT, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_run_font(prefix_run, bold=True)
        body_run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(body_run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        run = paragraph.add_run(line)
        set_run_font(run, CODE_FONT, size=9)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, align_center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_borders(table)
    table.autofit = True

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, size=font_size, align_center=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size, align_center=True)

    doc.add_paragraph()
    return table


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("운영체제 (Operating Systems)\nTerm Project 최종 보고서")
    set_run_font(run, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("공항 체크인 카운터 스케줄러")
    set_run_font(run, size=18, bold=True)

    add_paragraph(doc)
    add_paragraph(doc, "본 보고서는 공항 체크인 카운터 환경을 운영체제의 비선점형 CPU 스케줄링 문제로 모델링하고, 여러 스케줄링 정책의 Average Turnaround Time(ATT)을 비교 분석한 결과를 정리한 것이다.")

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["팀 명", TEAM_NAME],
            ["팀 원", f"{MEMBER_NAME} / {MEMBER_ID}"],
            ["", ""],
            ["", ""],
        ],
        font_size=10,
    )

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "목 차", level=1)
    toc_items = [
        "1. 설계 개요",
        "  1.1 큐 아키텍처 설계",
        "  1.2 알고리즘 조합 및 근거",
        "  1.3 카운터 배정 전략",
        "2. 구현",
        "  2.1 시스템 구조",
        "  2.2 핵심 모듈 설명",
        "  2.3 실행 방법",
        "3. 시뮬레이션 결과",
        "  3.1 승객별 결과",
        "  3.2 등급별 평균 Turnaround Time",
        "  3.3 카운터별 통계",
        "4. Baseline 비교 분석",
        "  4.1 ATT 비교 표",
        "  4.2 비교 분석",
        "5. Trade-off 분석 및 한계",
        "6. 역할 분담 및 기여도",
        "7. 생성형 AI 활용 경험",
        "8. 결론",
    ]
    for item in toc_items:
        add_paragraph(doc, item)
    doc.add_page_break()


def add_design_sections(doc: Document) -> None:
    add_heading(doc, "1. 설계 개요", level=1)
    add_heading(doc, "1.1 큐 아키텍처 설계", level=2)
    add_paragraph(doc, "본 프로젝트에서는 승객 등급을 기준으로 한 하이브리드 Multi-Level Queue 구조를 사용하였다. 전체 ready queue에 도착한 승객을 저장한 뒤, 스케줄러가 선택 시점마다 승객 등급에 따라 First Queue, Business Queue, Economy Queue로 논리적으로 분리한다.")
    add_code_block(
        doc,
        """
Ready Queue
  ├─ First Queue
  ├─ Business Queue
  └─ Economy Queue
        """,
    )
    add_paragraph(doc, "카운터는 총 5개이며, C1은 First 우선, C2는 Business 우선, C3는 Economy 우선 카운터로 설정하였다. C4와 C5는 Flex 카운터로 두어 모든 등급의 승객을 처리할 수 있도록 하였다.")
    add_paragraph(doc, "전용 카운터는 자기 등급 승객이 대기 중이면 해당 등급 큐에서 먼저 승객을 선택한다. 단, 자기 등급 승객이 없는 경우 카운터를 idle 상태로 두면 전체 ATT가 증가할 수 있으므로 다른 등급 승객도 처리할 수 있도록 허용하였다. Flex 카운터는 등급 제한 없이 전체 ready queue에서 가장 높은 스케줄링 점수를 가진 승객을 선택한다.")
    add_paragraph(doc, "큐 간 관계는 고정 우선순위만 사용하는 방식이 아니라, 등급별 가중치와 대기 시간 증가 효과를 함께 반영하는 방식으로 설계하였다. 이를 통해 First와 Business 승객에게 일정한 우선권을 주면서도, Economy 승객이 오래 기다릴 경우 선택될 수 있도록 하였다.")

    add_heading(doc, "1.2 알고리즘 조합 및 근거", level=2)
    add_paragraph(doc, "본 스케줄러는 하나의 알고리즘만 사용하지 않고 Multi-Level Queue, Priority Scheduling, HRRN/Aging, SJF tie-break, FCFS tie-break를 조합하였다.")
    add_table(
        doc,
        ["적용 위치", "알고리즘", "선택 근거", "기대 효과"],
        [
            ["등급별 대기열 전체", "Multi-Level Queue", "승객 등급이 First, Business, Economy로 구분되므로 등급별 큐 구조가 문제 상황에 적합하다.", "등급별 정책 적용과 카운터 배정 기준을 명확히 할 수 있다."],
            ["승객 등급 우선순위", "Priority Scheduling", "항공 체크인에서는 First와 Business 승객에게 상대적으로 높은 우선순위를 부여할 필요가 있다.", "First/Business 승객의 과도한 대기를 줄이고 등급 기반 서비스 정책을 반영한다."],
            ["전체 승객 선택 점수", "HRRN / Aging", "고정 우선순위만 사용하면 Economy 승객이 오래 대기할 수 있으므로 대기 시간이 길수록 우선순위가 증가하도록 한다.", "Starvation을 완화하고 오래 기다린 승객도 선택될 수 있게 한다."],
            ["점수 동률 처리", "Non-preemptive SJF", "같은 조건에서는 service_time이 짧은 승객을 먼저 처리하는 것이 평균 Turnaround Time 감소에 유리하다.", "짧은 작업을 빠르게 완료하여 전체 ATT 감소를 기대할 수 있다."],
            ["최종 동률 처리", "FCFS", "점수와 service_time이 모두 같을 경우 먼저 도착한 승객을 먼저 처리하는 것이 공정하다.", "동일 조건에서 처리 순서를 안정적으로 결정하고 공정성을 보장한다."],
        ],
        font_size=8,
    )
    add_paragraph(doc, "우리 팀 스케줄러의 선택 점수는 다음 식으로 계산하였다.")
    add_code_block(
        doc,
        """
waiting_time = current_time - arrival_time
response_ratio = (waiting_time + service_time) / service_time
score = response_ratio * class_weight

class_weight:
  First    = 1.5
  Business = 1.1
  Economy  = 1.0
        """,
    )
    add_paragraph(doc, "승객 선택 기준은 score가 높은 승객, service_time이 짧은 승객, arrival_time이 빠른 승객, passenger_id가 작은 승객 순서로 적용된다.")

    add_heading(doc, "1.3 카운터 배정 전략", level=2)
    add_paragraph(doc, "카운터는 과제 조건에 따라 총 5개를 사용하였다.")
    add_table(
        doc,
        ["카운터", "유형", "배정 전략"],
        [
            ["C1", "First 우선", "First 승객이 있으면 First Queue에서 먼저 선택"],
            ["C2", "Business 우선", "Business 승객이 있으면 Business Queue에서 먼저 선택"],
            ["C3", "Economy 우선", "Economy 승객이 있으면 Economy Queue에서 먼저 선택"],
            ["C4", "Flex", "전체 ready queue에서 Weighted HRRN 점수가 가장 높은 승객 선택"],
            ["C5", "Flex", "전체 ready queue에서 Weighted HRRN 점수가 가장 높은 승객 선택"],
        ],
    )
    add_paragraph(doc, "전용 카운터인 C1, C2, C3는 기본적으로 자기 등급 승객을 우선 처리한다. 그러나 자기 등급 승객이 없는 경우에는 다른 등급 승객도 처리하도록 허용하였다. 본 과제의 성능 지표가 ATT 하나이므로 카운터를 비워 두는 것보다 다른 승객을 처리하여 idle time을 줄이는 것이 전체 성능에 유리하다고 판단하였다.")
    add_paragraph(doc, "Flex 카운터인 C4와 C5는 특정 등급에 고정하지 않고 모든 ready queue 승객을 후보로 본다. 이때 단순 고정 우선순위가 아니라 Weighted HRRN 점수를 기준으로 선택하여 높은 등급 승객의 우선권과 오래 기다린 승객의 공정성을 함께 고려하였다.")


def add_implementation_sections(doc: Document) -> None:
    add_heading(doc, "2. 구현", level=1)
    add_heading(doc, "2.1 시스템 구조", level=2)
    add_paragraph(doc, "본 프로젝트는 Python으로 구현하였으며, 입력 처리, 시뮬레이션 엔진, 스케줄러 전략, 결과 출력 모듈을 분리하여 구성하였다. 전체 구조는 Strategy Pattern을 기반으로 하며, 동일한 SimulationEngine에서 스케줄러 전략만 교체하여 Baseline과 우리 팀 스케줄러를 비교할 수 있도록 설계하였다.")
    add_code_block(
        doc,
        """
input.txt
  ↓
scheduler.py
  ↓
parse_input_file()
  ↓
Passenger / Counter 객체 생성
  ↓
SimulationEngine
  ↓
SchedulerStrategy 선택
  ├─ Baseline A: FCFS
  ├─ Baseline B: Priority
  ├─ Baseline C: SJF
  └─ OurScheduler
  ↓
SimulationResult 생성
  ↓
CSV / Log / Graph 출력
        """,
    )
    add_table(
        doc,
        ["파일명", "역할"],
        [
            ["models.py", "승객, 카운터, 시뮬레이션 결과 데이터 모델 정의"],
            ["simulation.py", "이산 사건 기반 시뮬레이션 엔진 구현"],
            ["strategies.py", "Baseline 스케줄러와 우리 팀 스케줄러 구현"],
            ["scheduler.py", "입력 파싱, 실행 옵션 처리, 전체 실행 흐름 제어"],
            ["report_utils.py", "ATT 비교 그래프 생성 및 결과 파일 출력 보조"],
        ],
    )
    add_paragraph(doc, "이 구조를 통해 시뮬레이션 로직과 스케줄링 정책을 분리하였다. 새로운 스케줄링 알고리즘을 추가하더라도 SimulationEngine을 수정하지 않고 SchedulerStrategy를 상속한 클래스를 추가하는 방식으로 확장할 수 있다.")

    add_heading(doc, "2.2 핵심 모듈 설명", level=2)
    add_paragraph(doc, "핵심 모듈은 SimulationEngine과 SchedulerStrategy이다. SimulationEngine은 현재 시간 기준으로 도착한 승객을 ready queue에 추가하고, 서비스가 완료된 승객을 처리한 뒤, idle 상태인 카운터에 대해 스케줄러에게 다음 승객 선택을 요청한다.")
    add_paragraph(doc, "승객이 한 번 카운터에 배정되면 completion_time까지 current_passenger로 유지되며, 중간에 다른 승객으로 교체되지 않는다. 따라서 구현은 과제에서 요구한 비선점형 조건을 만족한다.")
    add_code_block(
        doc,
        """
while 모든 승객이 완료되지 않았으면:
    현재 시간까지 도착한 승객을 ready queue에 추가한다.
    완료 시간이 된 카운터의 승객을 completion 처리한다.

    idle 상태인 카운터가 있으면:
        scheduler.select_next_passenger()를 호출한다.
        선택된 승객을 해당 카운터에 배정한다.
        service_start_time과 completion_time을 기록한다.

    다음 이벤트 시간으로 이동한다.
        다음 이벤트 = 다음 승객 도착 시간 또는 다음 서비스 완료 시간
        """,
    )
    add_paragraph(doc, "SchedulerStrategy는 스케줄러의 공통 인터페이스이다. 모든 스케줄러는 select_next_passenger(ready_queue, counters, current_time, counter)를 구현하며, 이 함수는 현재 ready queue에서 다음에 처리할 승객 1명을 반환한다.")
    add_table(
        doc,
        ["스케줄러", "설명"],
        [
            ["BaselineA_FCFS", "도착 시간이 빠른 승객부터 처리"],
            ["BaselineB_Priority", "First > Business > Economy 순서로 처리"],
            ["BaselineC_SJF", "service_time이 짧은 승객부터 처리"],
            ["OurScheduler", "Multi-Level Queue + Weighted HRRN + SJF + FCFS 조합"],
        ],
    )

    add_heading(doc, "2.3 실행 방법", level=2)
    add_paragraph(doc, "프로그램은 명령 프롬프트 또는 PowerShell에서 실행할 수 있다. 프로젝트 폴더로 이동한 뒤 다음 명령어를 실행한다.")
    add_code_block(
        doc,
        """
cd C:\\OS_WORKSPACE
python scheduler.py input.txt --scheduler all
        """,
    )
    add_paragraph(doc, "개별 스케줄러만 실행하려면 다음과 같이 입력한다.")
    add_code_block(
        doc,
        """
python scheduler.py input.txt --scheduler fcfs
python scheduler.py input.txt --scheduler priority
python scheduler.py input.txt --scheduler sjf
python scheduler.py input.txt --scheduler ours
        """,
    )
    add_paragraph(doc, "출력 폴더를 직접 지정할 수도 있다.")
    add_code_block(doc, "python scheduler.py input.txt --scheduler all --output output")
    add_table(
        doc,
        ["파일명", "내용"],
        [
            ["passenger_results.csv", "승객별 arrival, start, completion, turnaround 결과"],
            ["class_summary.csv", "등급별 평균 Turnaround Time"],
            ["counter_summary.csv", "카운터별 처리 승객 수, 총 처리 시간, idle time"],
            ["att_comparison.csv", "스케줄러별 ATT 비교표"],
            ["att_comparison.png", "ATT 비교 막대그래프"],
            ["simulation_log.txt", "시간별 시뮬레이션 로그"],
        ],
    )
    add_paragraph(doc, "프로그램 검증은 다음 명령어로 수행하였다.")
    add_code_block(doc, "python -m unittest -v")


def add_results_sections(doc: Document) -> None:
    passengers = read_csv(OUTPUT_DIR / "passenger_results.csv")
    class_summary = read_csv(OUTPUT_DIR / "class_summary.csv")
    counter_summary = read_csv(OUTPUT_DIR / "counter_summary.csv")

    add_heading(doc, "3. 시뮬레이션 결과", level=1)
    add_paragraph(doc, "본 절의 결과는 우리 팀 스케줄러(Our Scheduler: MLQ + Weighted HRRN + SJF)를 기본 스케줄러로 실행한 결과이다.")
    add_paragraph(doc, "참고로 과제 PDF의 통계 요약에는 전체 service_time 합이 374로 표시되어 있으나, 실제 제공된 input.txt 기준 합계는 379이다. 따라서 본 구현과 검증은 실제 입력 파일의 값을 기준으로 수행하였다.")

    add_heading(doc, "3.1 승객별 결과 (기본 데이터셋)", level=2)
    rows = [
        [
            row["passenger_id"],
            row["class"],
            row["class_name"],
            row["arrival_time"],
            row["service_time"],
            row["service_start_time"],
            row["completion_time"],
            row["turnaround_time"],
            row["assigned_counter_id"],
        ]
        for row in passengers
    ]
    add_table(
        doc,
        ["ID", "등급", "등급명", "arrival", "service", "start", "completion", "turnaround", "counter"],
        rows,
        font_size=7,
    )

    add_heading(doc, "3.2 등급별 평균 Turnaround Time", level=2)
    add_table(
        doc,
        ["등급", "등급명", "승객 수", "평균 Turnaround Time"],
        [[row["class"], row["class_name"], row["passenger_count"], row["average_turnaround_time"]] for row in class_summary],
    )
    add_paragraph(doc, "우리 팀 스케줄러의 전체 ATT는 19.18이다. 등급별 평균 Turnaround Time은 First 20.38, Business 21.73, Economy 17.97로 나타났다.")

    add_heading(doc, "3.3 카운터별 통계", level=2)
    add_table(
        doc,
        ["카운터", "유형", "처리 승객 수", "총 처리 시간", "유휴 시간", "처리 승객 ID"],
        [
            [
                row["counter_id"],
                row["counter_type"],
                row["processed_count"],
                row["total_service_time"],
                row["idle_time"],
                row["processed_passenger_ids"],
            ]
            for row in counter_summary
        ],
        font_size=8,
    )
    add_paragraph(doc, "카운터별 총 처리 시간의 합은 379로, 실제 input.txt의 전체 service_time 합과 일치한다. 이는 모든 승객이 누락 없이 처리되었음을 의미한다.")


def add_comparison_sections(doc: Document) -> None:
    att_comparison = read_csv(OUTPUT_DIR / "att_comparison.csv")

    add_heading(doc, "4. Baseline 비교 분석", level=1)
    add_heading(doc, "4.1 ATT 비교 표", level=2)
    add_table(
        doc,
        ["스케줄러", "ATT", "우리 스케줄러의 해당 기준 대비 개선율(%)"],
        [[row["scheduler_name"], row["ATT"], row["improvement_rate"]] for row in att_comparison],
    )

    graph_path = OUTPUT_DIR / "att_comparison.png"
    if graph_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(graph_path), width=Inches(5.8))
        caption = doc.add_paragraph("그림 1. 스케줄러별 Average Turnaround Time 비교")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(caption.runs[0], size=9)

    add_heading(doc, "4.2 비교 분석", level=2)
    add_paragraph(doc, "Baseline A인 FCFS의 ATT는 19.52이고, 우리 팀 스케줄러의 ATT는 19.18이다. 우리 스케줄러는 FCFS 대비 약 1.74% 개선되었다. FCFS는 도착 순서만 고려하기 때문에 긴 service_time 승객이 앞에 있을 경우 뒤 승객들이 함께 지연되는 Convoy Effect가 발생할 수 있다. 반면 우리 스케줄러는 HRRN 점수와 service_time tie-break를 사용하여 일부 짧은 작업을 더 빠르게 처리할 수 있었다.")
    add_paragraph(doc, "Baseline B인 고정 우선순위 방식의 ATT는 22.50으로 가장 나빴다. First와 Business 승객에게 매우 유리하지만 Economy 승객이 크게 밀리기 때문에 전체 평균 Turnaround Time이 증가하였다. 우리 스케줄러는 등급별 가중치를 사용하되 대기 시간 증가 효과를 함께 반영하여 Economy 승객이 계속 뒤로 밀리는 문제를 완화하였다.")
    add_paragraph(doc, "Baseline C인 Non-preemptive SJF의 ATT는 14.90으로 전체 ATT만 보면 가장 우수하였다. 이는 SJF가 짧은 service_time 승객을 먼저 처리하여 평균 완료 시간을 줄이는 데 강하기 때문이다. 그러나 SJF는 등급을 고려하지 않기 때문에 service_time이 긴 First 승객의 Turnaround Time이 크게 증가하였다. 본 과제에서 ATT가 유일한 지표이기는 하지만, 항공 체크인 서비스의 등급 우선 정책을 고려하면 순수 SJF는 서비스 정책 측면에서 한계가 있다.")
    add_paragraph(doc, "우리 팀 스케줄러는 순수 SJF보다 전체 ATT는 높지만, SJF에서 크게 악화된 First 승객의 Turnaround Time을 줄이고 등급별 결과가 한쪽으로 과도하게 치우치지 않도록 설계하였다. 다만 Business 평균 Turnaround Time은 FCFS와 SJF보다 높게 나타났으므로, 이 부분은 본 스케줄러의 한계로 볼 수 있다.")


def add_tradeoff_section(doc: Document) -> None:
    scheduler_class_files = {
        "FCFS": OUTPUT_DIR / "fcfs_class_summary.csv",
        "Priority": OUTPUT_DIR / "priority_class_summary.csv",
        "SJF": OUTPUT_DIR / "sjf_class_summary.csv",
        "Ours": OUTPUT_DIR / "ours_class_summary.csv",
    }
    rows = []
    for scheduler_name, path in scheduler_class_files.items():
        summary = {row["class_name"]: row["average_turnaround_time"] for row in read_csv(path)}
        rows.append([
            scheduler_name,
            summary.get("First", ""),
            summary.get("Business", ""),
            summary.get("Economy", ""),
        ])

    add_heading(doc, "5. Trade-off 분석 및 한계", level=1)
    add_table(doc, ["스케줄러", "First 평균 TAT", "Business 평균 TAT", "Economy 평균 TAT"], rows)
    add_paragraph(doc, "SJF는 전체 ATT가 14.90으로 가장 낮지만 First 평균 Turnaround Time이 44.25로 매우 높게 나타났다. 이는 service_time이 긴 First 승객들이 짧은 작업 뒤로 계속 밀렸기 때문이다. 즉, SJF는 평균 성능에는 강하지만 등급 우선 정책과 공정성 측면에서는 약점이 있다.")
    add_paragraph(doc, "Priority 방식은 First 13.25, Business 11.09로 상위 등급 승객에게 유리하지만 Economy 평균 Turnaround Time이 28.94까지 증가하였다. 고정 우선순위만 사용할 경우 낮은 등급 승객의 starvation 위험이 커진다는 한계를 보여준다.")
    add_paragraph(doc, "우리 팀 스케줄러는 First 20.38, Business 21.73, Economy 17.97로 특정 등급 하나에만 성능이 과도하게 집중되지 않도록 설계되었다. 다만 Business 평균 Turnaround Time이 FCFS보다 약간 높아졌고, 순수 SJF보다 전체 ATT가 높다는 한계가 있다.")
    add_paragraph(doc, "결론적으로 우리 스케줄러는 최저 ATT만을 달성하는 알고리즘은 아니지만, 등급별 가중치와 HRRN aging을 함께 사용하여 상위 등급 우선 정책과 낮은 등급 승객의 starvation 방지를 동시에 고려하였다. 이는 항공 체크인이라는 실제 서비스 환경에 더 적합한 절충안이라고 판단하였다.")


def add_remaining_sections(doc: Document) -> None:
    add_heading(doc, "6. 역할 분담 및 기여도", level=1)
    add_table(
        doc,
        ["이름", "학번", "담당 역할", "기여도(%)"],
        [
            [MEMBER_NAME, MEMBER_ID, "", ""],
            ["", "", "", ""],
            ["", "", "", ""],
            ["", "", "", ""],
        ],
    )
    add_paragraph(doc, "전체 코드 구조와 실행 흐름을 검토하였으며, 최종 결과에 대해 실행 검증을 수행하였다.")

    add_heading(doc, "7. 생성형 AI 활용 경험", level=1)
    add_paragraph(doc, "프로젝트 수행 과정에서 생성형 AI를 설계 정리, 구현 체크리스트 작성, CPM 네트워크 구성, 코드 구조 설계, 테스트 실패 원인 분석, 최종보고서 초안 작성에 활용하였다.")
    add_paragraph(doc, "사용한 주요 프롬프트 및 작업 요청 예시는 다음과 같다.")
    add_code_block(
        doc,
        """
현재 프로젝트에 들어있는 문서 3가지를 보고 어떤 식으로 과제를 진행해야 할지 알려줘.
implementation_checklist.txt를 기반으로 CPM network를 구현해줘.
4차: 전체 통합 후 python scheduler.py input.txt --scheduler all 실행 이 부분 진행해줘.
최종 보고서의 1.1, 1.2, 1.3에 들어갈 내용을 작성해줘.
        """,
    )
    add_paragraph(doc, "생성형 AI를 활용하면서 과제 요구사항을 빠르게 구조화할 수 있었고, Strategy Pattern 기반의 구현 방향을 명확히 정할 수 있었다. 또한 테스트 실패 원인을 분석하는 과정에서 PDF의 통계 요약과 실제 input.txt 데이터가 일치하지 않는 문제를 발견할 수 있었다.")
    add_paragraph(doc, "다만 생성형 AI가 제안한 수치나 구조를 그대로 신뢰할 수는 없었다. 실제 입력 파일을 기준으로 service_time 합계, 승객 수, ATT 계산 결과를 직접 검증해야 했으며, 최종 보고서에는 실제 실행 결과를 기준으로 수치를 반영하였다.")

    add_heading(doc, "8. 결론", level=1)
    add_paragraph(doc, "본 프로젝트에서는 공항 체크인 카운터 문제를 운영체제의 비선점형 CPU 스케줄링 문제로 모델링하고, 여러 스케줄링 정책을 직접 구현하여 비교하였다. Baseline으로 FCFS, 고정 우선순위, Non-preemptive SJF를 구현하였고, 우리 팀 스케줄러는 Multi-Level Queue, Priority Weight, HRRN/Aging, SJF, FCFS를 조합하여 설계하였다.")
    add_paragraph(doc, "실험 결과, 우리 팀 스케줄러의 ATT는 19.18로 FCFS와 고정 우선순위 방식보다 개선되었다. 반면 순수 SJF는 ATT 14.90으로 가장 낮았지만, First 승객의 평균 Turnaround Time이 크게 증가하는 문제가 있었다. 이를 통해 평균 성능만을 최적화하는 알고리즘과 서비스 등급 정책을 함께 고려하는 알고리즘 사이에는 명확한 trade-off가 있음을 확인하였다.")
    add_paragraph(doc, "본 프로젝트를 통해 CPU 스케줄링 알고리즘이 단순히 이론적인 평균 시간 최소화 문제에 그치지 않고, 실제 서비스 정책과 공정성 조건에 따라 다르게 설계되어야 함을 이해할 수 있었다. 향후 개선 방향으로는 class_weight 값을 자동 탐색하거나, 시간대별 승객 분포에 따라 Flex 카운터의 우선순위를 동적으로 조정하는 방식을 고려할 수 있다.")


def main() -> None:
    template = find_template()
    doc = Document(str(template)) if template else Document()
    clear_document_body(doc)
    set_default_styles(doc)

    add_title_page(doc)
    add_toc(doc)
    add_design_sections(doc)
    add_implementation_sections(doc)
    add_results_sections(doc)
    add_comparison_sections(doc)
    add_tradeoff_section(doc)
    add_remaining_sections(doc)

    try:
        doc.save(REPORT_PATH)
        print(REPORT_PATH)
    except PermissionError:
        doc.save(REVIEWED_REPORT_PATH)
        print(REVIEWED_REPORT_PATH)


if __name__ == "__main__":
    main()
